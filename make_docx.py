#!/usr/bin/env python3
"""Render a pharmacology lecture's entries.json into a bilingual Word document
with embedded frames and clickable jump-back timestamp links.

Typography is fully unified and explicit (no reliance on Word's built-in styles):
English -> Times New Roman, Chinese -> KaiTi (楷体). Body text is one size/weight,
headings are one size/weight. Adjust the constants below to taste.

Usage: python3 make_docx.py <entries.json> [output.docx]
"""
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx not installed. Run: pip3 install python-docx")

# ----- TYPOGRAPHY (edit here to change the whole document) -------------------
FONT_EN = "Times New Roman"   # English / Latin font
FONT_CN = "KaiTi"             # Chinese font (楷体). On some Macs use "STKaiti".
SIZE_TITLE = 16               # ALL headings (title + sections) one size, pt
SIZE_H1 = 16                  # same as title -> headings fully unified
SIZE_BODY = 12                # ALL body text (points, tags, links), pt
COLOR_HEADING = RGBColor(0x0B, 0x5A, 0x5A)
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)
COLOR_LINK = RGBColor(0x1A, 0x5F, 0xB4)
COLOR_TAG = RGBColor(0x59, 0x59, 0x59)
# -----------------------------------------------------------------------------

CATEGORIES = [
    ("mechanism", "作用机制", "Mechanism of Action"),
    ("PK", "药代动力学", "Pharmacokinetics (PK)"),
    ("PD", "药效动力学", "Pharmacodynamics (PD)"),
    ("SAR", "构效关系", "Structure–Activity Relationship"),
    ("indication", "适应症", "Indications"),
    ("contraindication", "禁忌症", "Contraindications"),
    ("ADR", "不良反应", "Adverse Drug Reactions"),
    ("interaction", "药物相互作用", "Drug Interactions"),
    ("concept", "其他要点", "Other Key Points"),
]


def _apply_fonts(rpr):
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)


def run(paragraph, text, size=SIZE_BODY, bold=False, color=COLOR_BODY):
    r = paragraph.add_run(text)
    r.font.name = FONT_EN
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    _apply_fonts(r._element.get_or_add_rPr())
    return r


def set_style_font(doc, style_name, size, bold):
    try:
        st = doc.styles[style_name]
    except KeyError:
        return
    st.font.name = FONT_EN
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = COLOR_BODY
    _apply_fonts(st.element.get_or_add_rPr())


def heading(doc, text, size=SIZE_H1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run(p, text, size=size, bold=True, color=COLOR_HEADING)
    return p


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_EN); rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_CN); rPr.append(rfonts)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1A5FB4"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(SIZE_BODY * 2)); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def jump_url(url, ts_seconds):
    if not url or ts_seconds is None:
        return url or ""
    sep = "&" if "?" in url else "?"
    return f"{url}{sep}t={int(ts_seconds)}s"


def tag_line(e):
    bits = []
    for key, label in (("drugs", "药物"), ("targets", "靶点"),
                       ("pathways", "通路"), ("topics", "主题")):
        vals = e.get(key) or []
        if vals:
            bits.append(f"{label}: {', '.join(vals)}")
    return "  |  ".join(bits)


def main():
    if len(sys.argv) < 2:
        sys.exit("Usage: python3 make_docx.py <entries.json> [output.docx]")
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else src.rsplit(".", 1)[0] + ".docx"
    with open(src, encoding="utf-8") as f:
        data = json.load(f)
    v = data.get("video", {})
    entries = data.get("entries", [])

    doc = Document()
    set_style_font(doc, "Normal", SIZE_BODY, False)
    set_style_font(doc, "List Bullet", SIZE_BODY, False)

    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(2)
    run(tp, v.get("title", "Pharmacology Notes"), size=SIZE_TITLE, bold=True, color=COLOR_HEADING)
    mbits = [x for x in (v.get("course"), v.get("duration"), v.get("url")) if x]
    if mbits:
        mp = doc.add_paragraph()
        run(mp, "  ·  ".join(mbits), size=SIZE_BODY, color=COLOR_TAG)

    if data.get("overview_zh"):
        heading(doc, "概述 Overview")
        p = doc.add_paragraph()
        run(p, data["overview_zh"])

    by_cat = {}
    for e in entries:
        by_cat.setdefault(e.get("category", "concept"), []).append(e)

    def render_entry(e):
        p = doc.add_paragraph(style="List Bullet")
        term = e.get("term_en", "")
        if e.get("term_zh"):
            term = f"{term} / {e['term_zh']}" if term else e["term_zh"]
        run(p, term, bold=False)  # body weight uniform (no bold in body)
        if e.get("explanation_zh"):
            run(p, " — " + e["explanation_zh"])
        tl = tag_line(e)
        if tl:
            tpp = doc.add_paragraph()
            tpp.paragraph_format.left_indent = Inches(0.4)
            run(tpp, tl, color=COLOR_TAG)
        ju = jump_url(v.get("url"), e.get("ts_seconds"))
        if ju:
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Inches(0.4)
            add_hyperlink(lp, ju, f"▶ 回看 Re-watch @ {e.get('ts_hhmmss','')}")
        fr = e.get("frame")
        if fr and os.path.exists(fr):
            try:
                doc.add_picture(fr, width=Inches(4.5))
            except Exception:
                pass

    for cat, zh, en in CATEGORIES:
        items = by_cat.get(cat)
        if not items:
            continue
        heading(doc, f"{zh}  {en}")
        for e in items:
            render_entry(e)

    known = {c for c, _, _ in CATEGORIES}
    other = [e for e in entries if e.get("category", "concept") not in known]
    if other:
        heading(doc, "未分类 Uncategorized")
        for e in other:
            render_entry(e)

    doc.save(out)
    print(f"Wrote Word document: {out}  ({len(entries)} entries)")


if __name__ == "__main__":
    main()
